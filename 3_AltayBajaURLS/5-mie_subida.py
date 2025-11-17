import time
import csv
import os
import requests
from docx import Document
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.common.exceptions import ElementClickInterceptedException, TimeoutException, NoSuchElementException

# Importar configuración
try:
    import config
except ImportError:
    print("ERROR: No se encontró el archivo config.py")
    print("Por favor, copia config.example.py como config.py y completa las credenciales.")
    exit(1)

# Configura las opciones del navegador
options = Options()
if config.HEADLESS_MODE:
    options.add_argument("--headless")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

# Usa WebDriver Manager para gestionar el driver automáticamente
# Soporta Edge, Chrome y Firefox según configuración
if config.BROWSER.lower() == 'edge':
    driver = webdriver.Edge(service=Service(EdgeChromiumDriverManager().install()), options=options)
elif config.BROWSER.lower() == 'chrome':
    from webdriver_manager.chrome import ChromeDriverManager
    from selenium.webdriver.chrome.service import Service as ChromeService
    driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
elif config.BROWSER.lower() == 'firefox':
    from webdriver_manager.firefox import GeckoDriverManager
    from selenium.webdriver.firefox.service import Service as FirefoxService
    driver = webdriver.Firefox(service=FirefoxService(GeckoDriverManager().install()), options=options)
else:
    raise ValueError(f"Navegador no soportado: {config.BROWSER}. Usa 'edge', 'chrome' o 'firefox'")

# Tamaño de ventana específico para asegurar que los elementos necesarios sean visibles
driver.set_window_size(config.WINDOW_SIZE[0], config.WINDOW_SIZE[1])

# URLs y configuración desde archivo de configuración
url = config.WORDPRESS_LOGIN_URL
API_key = config.GOOGLE_MAPS_API_KEY
carpeta_word = config.GENERATED_TEMPLATES_FOLDER

try:
    # Abre la página
    driver.get(url)

    # Espera a que los campos de usuario y contraseña estén visibles
    username_field = WebDriverWait(driver, 20).until(
        EC.visibility_of_element_located((By.NAME, 'log'))
    )
    password_field = WebDriverWait(driver, 20).until(
        EC.visibility_of_element_located((By.NAME, 'pwd'))
    )

    # Establece el valor usando JavaScript (desde configuración)
    driver.execute_script("arguments[0].value = arguments[1];", username_field, config.WORDPRESS_USERNAME)
    driver.execute_script("arguments[0].value = arguments[1];", password_field, config.WORDPRESS_PASSWORD)

    # Clic en el botón de inicio de sesión
    login_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.ID, 'wp-submit'))
    )
    login_button.click()

    # Espera a que la URL cambie después del primer login
    WebDriverWait(driver, 10).until(EC.url_changes(url))

    # def second_login(username, password):
    #     try:
    #         username_field = WebDriverWait(driver, 10).until(
    #             EC.visibility_of_element_located((By.ID, 'user_login'))
    #         )
    #         password_field = WebDriverWait(driver, 10).until(
    #             EC.visibility_of_element_located((By.ID, 'user_pass'))
    #         )

    #         driver.execute_script("arguments[0].value = arguments[1];", username_field, username)
    #         driver.execute_script("arguments[0].value = arguments[1];", password_field, password)

    #         login_button = WebDriverWait(driver, 10).until(
    #             EC.element_to_be_clickable((By.ID, 'wp-submit'))
    #         )
    #         login_button.click()
    #     except TimeoutException:
    #         print("Tiempo de espera agotado durante el segundo inicio de sesión.")
    #     except Exception as e:
    #         print(f"Error durante el segundo inicio de sesión: {e}")

    # # Realiza el segundo inicio de sesión
    # second_login('gruposanz', 'pinkstone_01')

    # # Espera a que la URL cambie después del segundo login
    # WebDriverWait(driver, 10).until(EC.url_changes(url))

    # Función para leer el documento de Word
    def leer_documento(path):
        try:
            doc = Document(path)
            contenido = []
            for para in doc.paragraphs:
                if para.text.strip():
                    html_paragraph = para.text
                    for run in para.runs:
                        if run.bold:
                            html_paragraph = html_paragraph.replace(run.text, f"<strong>{run.text}</strong>")
                    contenido.append(html_paragraph)
            return contenido
        except Exception as e:
            print(f"Error al leer el documento {path}: {e}")
            return []

    # Función para listar todos los archivos .docx en la ruta y sus subcarpetas
    def listar_archivos_docx(ruta):
        archivos_word = []
        for root, dirs, files in os.walk(ruta):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~$'):
                    archivos_word.append(os.path.join(root, file))
        return archivos_word

    # Función para extraer la descripción
    def extraer_descripcion(contenido):
        extraido = []
        extraer = False

        for texto in contenido:
            if texto.strip().lower() == "descripción":
                extraer = True
                continue
            elif texto.strip().lower() == "etiqueta de producto título de la ficha":
                extraer = False
                break
            if extraer:
                extraido.append(texto)

        return extraido

    # Función para extraer el título
    def extraer_titulo(contenido):
        titulo = ""
        for i, texto in enumerate(contenido):
            if "Título de la ficha" in texto:
                if i + 1 < len(contenido):
                    titulo = contenido[i + 1].strip()
                    break
        return titulo
    
    # Función para generar el código de la sección "Mapa CP" en Wordpress
    def generar_iframe_gmaps(API_key, address):
    
        encoded_address = requests.utils.quote(address)
        url = f"https://www.google.com/maps/embed/v1/place?key={API_key}&q={encoded_address}"
        
        response = requests.get(url)
        if response.status_code == 200:
            iframe = f'<iframe src="{response.url}" width="600" height="450" style="border:0;" allowfullscreen="allowfullscreen"></iframe>'
            return iframe
        else:
            return ''

    # Función para extraer el párrafo que sigue a "Etiqueta de Producto Título de la ficha"
    def etiqueta_product(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "etiqueta de producto título de la ficha":
                buscar_posterior = True
        return ""
    
    # Función para extraer el párrafo que sigue a "Proveedor"
    def proveedor(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "proveedor en el sitio wordpress":
                buscar_posterior = True
        return ""
    
    # Función para extraer el párrafo que sigue a "Localizador Google Maps"
    def ubicacion(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "localizador google maps":
                buscar_posterior = True
        return ""

    # Función para extraer el párrafo que sigue a "Descripción Corta"
    def descripcion_corta(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "descripción corta":
                buscar_posterior = True
        return ""

    # Función para extraer el párrafo que sigue a "Título Seo"
    def titulo_seo(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "título seo":
                buscar_posterior = True
        return ""

    # Función para extraer el párrafo que sigue a "MetaDescription"
    def meta_description(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "metadescription":
                buscar_posterior = True
        return ""

    # Función para extraer el párrafo que sigue a "Frase Clave Objetivo"
    def frase_clave(doc):
        buscar_posterior = False
        for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "frase clave objetivo":
                buscar_posterior = True
        return ""


    # Función para convertir el contenido del documento en HTML simple, incluyendo la tabla
    def convertir_a_html(contenido):
        html = "<html><body>"
        for para in contenido:
            html += f"{para}"
        html += "</body></html>"
        return html

    # Lista los archivos de Word en la carpeta y sus subcarpetas, ignorando los archivos temporales
    archivos_word = listar_archivos_docx(carpeta_word)
    #archivos_word_sorted = sorted(archivos_word)  # Opcional: ordenar los archivos

    if not archivos_word:
        print(f"No se encontraron archivos .docx en la ruta: {carpeta_word}")
    else:
        print(f"Se encontraron {len(archivos_word)} archivos .docx para procesar.")

    # Abre el archivo CSV y lee las URLs
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    csv_path = os.path.join(SCRIPT_DIR, config.URLS_CSV_FILE.replace('./', ''))
    try:
        with open(csv_path, newline='', encoding='utf-8') as csvfile:
            url_reader = csv.reader(csvfile)

            for i, row in enumerate(url_reader):
                if i >= len(archivos_word):
                    print("No hay más archivos .docx para asignar a las URLs.")
                    break
                if not row:
                    print(f"Fila {i+1} vacía. Saltando...")
                    continue
                product_url = row[0].strip()
                if not product_url:
                    print(f"URL en la fila {i+1} está vacía. Saltando...")
                    continue

                time.sleep(10)
                archivo = archivos_word[i]
                print(f"Procesando URL: {product_url} con archivo: {archivo}")

                driver.get(product_url)

                    # Intentar hacer clic en "Editar producto"
                try:
                        edit_product_button = WebDriverWait(driver, 30).until(
                            EC.element_to_be_clickable((By.LINK_TEXT, 'Editar producto'))
                        )
                        edit_product_button.click()
                except TimeoutException:
                        print(f"No se encontró el botón 'Editar producto' en la URL: {product_url}. Saltando...")
                        continue
                except Exception as e:
                        print(f"Error al hacer clic en 'Editar producto' en la URL: {product_url}. Error: {e}. Saltando...")
                        continue

                # Extraer y procesar datos del archivo Word
                contenido = leer_documento(archivo)
                titulo = extraer_titulo(contenido)
                contenido_descripcion = extraer_descripcion(contenido)
                doc = Document(archivo)
                
                descripcion_html = convertir_a_html(contenido_descripcion)
                contenido_corto = descripcion_corta(doc)
                prov = proveedor(doc)
                etiq_prod = etiqueta_product(doc)
                tit_seo = titulo_seo(doc)
                meta_desc = meta_description(doc)
                frase_objetivo = frase_clave(doc)
                codigo_gmaps = generar_iframe_gmaps(API_key, ubicacion(doc))

                print(f"Título extraído: {titulo}")
                print(f"Descripción en HTML (primeros 100 caracteres): {descripcion_html[:100]}...")

                # Actualizar campos en WordPress
                try:
                        
                        # Actualizar título
                        title_field = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'title'))
                        )
                        title_field.clear()
                        title_field.send_keys(titulo)
                        print(f"Título enviado a WordPress: {titulo}")

                        # Se pasa al modo HTML de la descripción principal en caso de que esté en modo visual
                        html_content_button = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'content-html'))
                        )
                        html_content_button.click()

                        # Actualizar descripción principal

                        # Este modo de actualización usa la pestaña visual
                        # Dado que la plantilla actual usa HTML, se ha optado por el modo empleado más abajo 
                        # iframe = WebDriverWait(driver, 30).until(
                        #     EC.presence_of_element_located((By.ID, 'content_ifr'))
                        # )
                        # driver.switch_to.frame(iframe)
                        # print("Sección de descripción localizada.")

                        # editor_field = WebDriverWait(driver, 30).until(
                        #     EC.presence_of_element_located((By.ID, 'tinymce'))
                        # )
                        # editor_field.clear()
                        # driver.execute_script("arguments[0].innerHTML = arguments[1];", editor_field, descripcion_html)
                        # driver.switch_to.default_content()

                        textarea = WebDriverWait(driver, 30).until(
                            EC.presence_of_element_located((By.ID, 'content'))
                        )
                        textarea.clear()
                        driver.execute_script("arguments[0].value = arguments[1];", textarea, descripcion_html)
                        print("Descripción actualizada.")

                        # Marcar al proveedor y eliminar las marcas previas, si existen

                        prov_list = WebDriverWait(driver, 20).until(
                            EC.visibility_of_element_located((By.ID, 'proveedoreschecklist'))
                        )
                        prov_inputs = prov_list.find_elements(By.XPATH, "./li/label/input")

                        for prov_input in prov_inputs:
                            label = prov_input.find_element(By.XPATH, "parent::label")
                            text = label.text
                            if (prov not in text and prov_input.is_selected()) or (prov in text and not prov_input.is_selected()):
                                prov_input.click()
                        print("Proveedor actual seleccionado.")

                        # Actualizar frase clave objetivo
                        frase_field = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'focus-keyword-input-metabox'))
                        )
                        frase_field.clear()
                        frase_field.send_keys(Keys.CONTROL, 'a')
                        frase_field.send_keys(Keys.DELETE)
                        frase_field.send_keys(frase_objetivo)
                        print("Frase clave actualizada.")

                        # Eliminar etiquetas previas si existen
                        try:
                            WebDriverWait(driver, 5).until(
                                EC.visibility_of_element_located((By.CLASS_NAME, "tagchecklist"))
                            )

                            remove_buttons = driver.find_elements(By.XPATH, "//*[contains(@id, 'product_tag-check-num')]")
                            if remove_buttons:
                                while True:
                                    remove_buttons = driver.find_elements(By.XPATH, "//*[contains(@id, 'product_tag-check-num')]")
                                    if len(remove_buttons) > 0:
                                        button = remove_buttons[0]
                                        driver.execute_script("arguments[0].scrollIntoView(true);", button)
                                        driver.execute_script("arguments[0].click();", button)
                                    else:
                                        break
                                print("Etiquetas previas eliminadas.")
                        except TimeoutException:
                            print("No se encontró 'tagchecklist'. Esto es normal si no había etiquetas previas.")

                        # Actualizar etiqueta de producto
                        etiq_prodfield = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'new-tag-product_tag'))
                        )
                        etiq_prodfield.clear()
                        etiq_prodfield.send_keys(etiq_prod)
                        etiq_prodfield.send_keys(Keys.RETURN)
                        print("Etiqueta actualizada.")

                        # Actualizar título SEO
                        titseo_field = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'yoast-google-preview-title-metabox'))
                        )
    
                        titseo_field.send_keys(Keys.CONTROL, 'a')
                        titseo_field.send_keys(tit_seo)
                        print("Título SEO actualizado.")

                        # Actualizar meta descripción
                        metdesc_field = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, 'yoast-google-preview-description-metabox'))
                        )

                        metdesc_field.send_keys(Keys.CONTROL, 'a')
                        metdesc_field.send_keys(meta_desc)
                        print("Meta descripción actualizada.")

                        # Se pasa al modo visual de la descripción corta en caso de que esté en modo HTML
                        visual_excerpt_button = driver.find_element(By.ID, 'excerpt-tmce')
                        driver.execute_script("arguments[0].scrollIntoView(true);", visual_excerpt_button)
                        driver.execute_script("arguments[0].click();", visual_excerpt_button)

                        # Actualizar descripción corta
                        WebDriverWait(driver, 30).until(
                            EC.frame_to_be_available_and_switch_to_it((By.ID, "excerpt_ifr"))
                        )
                        short_description_body = WebDriverWait(driver, 30).until(
                            EC.visibility_of_element_located((By.ID, "tinymce"))
                        )

                        short_description_body.clear()
                        short_description_body.send_keys(contenido_corto)
                        print("Descripción corta actualizada.")

                        # Volver al contexto principal
                        driver.switch_to.default_content()

                        # Actualizar mapa
                        map_html_button = WebDriverWait(driver, 20).until(
                            EC.element_to_be_clickable((By.ID, "acf-editor-43-html"))
                        )
                        driver.execute_script("arguments[0].click();", map_html_button)

                        map_html_content = WebDriverWait(driver, 20).until(
                            EC.visibility_of_element_located((By.ID, "acf-editor-43"))
                        )
                        map_html_content.clear()
                        map_html_content.send_keys(codigo_gmaps)

                        print("Mapa actualizado.")

                        actualizar_button = WebDriverWait(driver, 30).until(
                            EC.presence_of_element_located((By.ID, "publish"))
                        )

                        # driver.execute_script("arguments[0].click();", actualizar_button)
                        # print("Contexto principal localizado.")
                        #Hace click en el botón después de desplazarse
                        #actualizar_button.click()
                        
                        driver.execute_script("arguments[0].scrollIntoView(true);", actualizar_button)
                        time.sleep(1)
                        driver.execute_script("arguments[0].click();", actualizar_button)
                        print("Se hizo clic en el botón de 'Actualizar'.")

                        # Esperar brevemente para asegurar que la actualización se complete
                        # print("Espera final activada.")
                        # time.sleep(5)

                except TimeoutException:
                        print(f"Tiempo de espera agotado al intentar actualizar la URL: {product_url}.")
                except Exception as e:
                        print(f"Error al actualizar el producto en la URL: {product_url}. Error: {e}")

    except FileNotFoundError:
        print(f"No se encontró el archivo CSV en la ruta: {csv_path}")
    except Exception as e:
        print(f"Error al abrir o procesar el archivo CSV: {e}")

finally:
    # Espera antes de cerrar el navegador
    time.sleep(20)  
    # Cierra el navegador
    driver.quit()
    print("Navegador cerrado.")
