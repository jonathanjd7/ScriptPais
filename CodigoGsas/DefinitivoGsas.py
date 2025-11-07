import time
import csv
import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException, TimeoutException, NoSuchElementException

# Obtener la ruta del directorio donde está el script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Configura las opciones de Edge
options = Options()
options.add_argument("--start-maximized")

# options.add_argument("--headless")  # Descomenta si deseas ejecutar sin interfaz gráfica
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

# Usa el driver desde el PATH del sistema (no requiere descarga)
# Asegúrate de tener msedgedriver.exe instalado o en el PATH
try:
    driver = webdriver.Edge(options=options)
except Exception as e:
    print(f"Error al inicializar Edge: {e}")
    print("Por favor, asegúrate de que Microsoft Edge WebDriver esté instalado.")
    print("Puedes descargarlo desde: https://developer.microsoft.com/en-us/microsoft-edge/tools/webdriver/")
    raise


# URL de la página de inicio de sesión
url = 'https://www.gsas.es/grupo-sanz-acceso/'

# Ruta a la carpeta con los documentos de Word (relativa al script)
# Busca la carpeta "Documentos" en el mismo directorio que el script
carpeta_word = os.path.join(script_dir, 'Documentos')
# Si no existe, intenta buscar en el directorio padre
if not os.path.exists(carpeta_word):
    carpeta_word = os.path.join(os.path.dirname(script_dir), 'Documentos')
# Si aún no existe, muestra un mensaje
if not os.path.exists(carpeta_word):
    print(f"ADVERTENCIA: No se encontró la carpeta 'Documentos' en {carpeta_word}")
    print("Por favor, asegúrate de que la carpeta 'Documentos' esté en el mismo directorio que el script.")
    print(f"Directorio del script: {script_dir}")

try:
    # Abre la página
    driver.get(url)
    

    # Espera a que los campos de usuario y contraseña estén visibles
    username_field = WebDriverWait(driver, 20).until(
        EC.visibility_of_element_located((By.NAME, 'log'))
    )
    password_field = WebDriverWait(driver, 20).until(
        EC.visibility_of_element_located((By.NAME, 'pwd'))
    )

    # Establece el valor usando JavaScript
    driver.execute_script("arguments[0].value = arguments[1];", username_field, 'gruposanz')
    driver.execute_script("arguments[0].value = arguments[1];", password_field, 'dM$NnF#fKwglDXOn3GYj&Et8')

    # Clic en el botón de inicio de sesión
    login_button = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, 'wp-submit'))
    )
    login_button.click()

    # Espera a que la URL cambie después del primer login
    WebDriverWait(driver, 20).until(EC.url_changes(url))

    def second_login(username, password):
        try:
            username_field = WebDriverWait(driver, 20).until(
                EC.visibility_of_element_located((By.ID, 'user_login'))
            )
            password_field = WebDriverWait(driver, 20).until(
                EC.visibility_of_element_located((By.ID, 'user_pass'))
            )

            driver.execute_script("arguments[0].value = arguments[1];", username_field, username)
            driver.execute_script("arguments[0].value = arguments[1];", password_field, password)

            login_button = WebDriverWait(driver, 20).until(
                EC.element_to_be_clickable((By.ID, 'wp-submit'))
            )
            login_button.click()
        except TimeoutException:
            print("Tiempo de espera agotado durante el segundo inicio de sesión.")
        except Exception as e:
            print(f"Error durante el segundo inicio de sesión: {e}")

    # Realiza el segundo inicio de sesión
    second_login('gruposanz', 'pinkstone_01')

    # Espera a que la URL cambie después del segundo login
    WebDriverWait(driver, 20).until(EC.url_changes(url))

    # Función para leer el documento de Word
    def leer_documento(path):
        try:
            doc = Document(path)
            contenido = []
            for para in doc.paragraphs:
                if para.text.strip():
                    html_paragraph = para.text
                    for run in para.runs:
                        if run.bold:
                            html_paragraph = html_paragraph.replace(run.text, f"<strong>{run.text}</strong>")
                    contenido.append(html_paragraph)
            return contenido
        except Exception as e:
            print(f"Error al leer el documento {path}: {e}")
            return []

    def _orden_archivo(path):
        """Extrae el número al inicio del nombre del archivo para ordenarlo correctamente"""
        nombre = os.path.splitext(os.path.basename(path))[0]
        # Extrae el número al inicio del nombre (antes del primer guion bajo o espacio)
        # Ejemplo: "1_Empresas_y_Tiendas..." -> 1
        try:
            # Intenta extraer el número al inicio
            numero_str = ""
            for char in nombre:
                if char.isdigit():
                    numero_str += char
                elif numero_str:  # Si ya encontramos dígitos y encontramos un no-dígito, paramos
                    break
            if numero_str:
                return int(numero_str)
            else:
                # Si no hay número, ordena alfabéticamente
                return float('inf')
        except (ValueError, IndexError):
            # Si no se puede convertir, ordena alfabéticamente al final
            return float('inf')

    # Función para listar todos los archivos .docx en la ruta y sus subcarpetas
    def listar_archivos_docx(ruta):
        archivos_word = []
        try:
            for root, dirs, files in os.walk(ruta):
                dirs.sort()
                files.sort()
                for file in files:
                    if file.endswith('.docx') and not file.startswith('~$'):
                        archivos_word.append(os.path.join(root, file))
        except Exception as e:
            print(f"Error al listar archivos en {ruta}: {e}")
        archivos_word.sort(key=_orden_archivo)
        return archivos_word

    # Función para extraer la descripción
    def extraer_descripcion(contenido):
        extraido = []
        extraer = False
        try:
         for texto in contenido:
            if texto.strip().lower() == "descripción":
                extraer = True
                continue
            elif texto.strip().lower() == "etiqueta de producto título de la ficha":
                extraer = False
                break
            if extraer:
                extraido.append(texto)
        except Exception as e:
         print(f"Error al extraer la descripción: {e}")
        return extraido

    # Función para extraer el título
    def extraer_titulo(contenido):
        titulo = ""
        try:
         for i, texto in enumerate(contenido):
            if "Título de la ficha" in texto:
                if i + 1 < len(contenido):
                    titulo = contenido[i + 1].strip()
                    break
        except Exception as e:
         print(f"Error al extraer el título: {e}")
        return titulo

    # Función para extraer la tabla de datos desde el documento Word
    def extraer_tabla(doc):
        tabla = []
        try:
         for table in doc.tables:
            for row in table.rows:
                fila = [cell.text.strip() for cell in row.cells]
                tabla.append(fila)
        except Exception as e:
         print(f"Error al extraer la tabla: {e}")
        return tabla

    # Función para extraer el párrafo que sigue a "Etiqueta de Producto Título de la ficha"
    def etiqueta_product(doc):
        buscar_posterior = False
        try:
         for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "etiqueta de producto título de la ficha":
                buscar_posterior = True
        except Exception as e:
         print(f"Error al extraer el párrafo posterior a 'Etiqueta de Producto Título de la ficha': {e}")
        return ""

    # Función para extraer el párrafo que sigue a "Descripción Corta"
    def descripcion_corta(doc):
        buscar_posterior = False
        try:
         for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "descripción corta":
                buscar_posterior = True
        except Exception as e:
         print(f"Error al extraer el párrafo posterior a 'Descripción Corta': {e}")
        return ""

    # Función para extraer el párrafo que sigue a "Título Seo"
    def titulo_seo(doc):
        buscar_posterior = False
        try:
         for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "título seo":
                buscar_posterior = True
        except Exception as e:
         print(f"Error al extraer el párrafo posterior a 'Título Seo': {e}")
        return ""

    # Función para extraer el párrafo que sigue a "MetaDescription"
    def meta_description(doc):
        buscar_posterior = False
        try:
         for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "metadescription":
                buscar_posterior = True
        except Exception as e:
         print(f"Error al extraer el párrafo posterior a 'MetaDescription': {e}")
        return ""

    # Función para extraer el párrafo que sigue a "Frase Clave Objetivo"
    def frase_clave(doc):
        buscar_posterior = False
        try:
         for para in doc.paragraphs:
            if buscar_posterior:
                return para.text.strip()
            if para.text.strip().lower() == "frase clave objetivo":
                buscar_posterior = True
        except Exception as e:
         print(f"Error al extraer el párrafo posterior a 'Frase Clave Objetivo': {e}")
        return ""

    # Función para extraer el precio que precede al símbolo €
    def extraer_precio(doc):
        try:
         for para in doc.paragraphs:
            texto = para.text.strip()
            if '€' in texto:
                partes = texto.split('€')
                if len(partes) > 0:
                    try:
                        precio_str = partes[0].strip().split()[-1]
                        precio_str = precio_str.replace('.', ',')
                        return precio_str
                    except IndexError:
                        print(f"No se pudo extraer el precio de la línea: {texto}")
        except Exception as e:
         print(f"Error al extraer el precio: {e}")
        return ""

    # Función para convertir el contenido del documento en HTML simple, incluyendo la tabla
    def convertir_a_html(contenido, tabla):
        html = "<html><body>"
        try:
         for para in contenido:
            para = para.replace(
                "bases de datos personalizados.", '<a href="https://gsas.es/bases-de-datos-de-empresas-a-medida/">bases de datos personalizados.</a>'
            ).replace(
                "PayPal", '<a href="https://www.paypal.com/">PayPal</a>'
            ).replace(
                "pedirnos información o pedir presupuesto", '<a href="https://gsas.es/contacto/">pedirnos información o pedir presupuesto</a>'
            ).replace(
                "campañas de envíos de email masivos.", '<a href="https://gsas.es/soluciones-todo-incluido-para-acciones-de-email-marketing/">campañas de envíos de email masivos.</a>'
            ).replace(
                "Solicita Fichero Demo Excel gratuito", '<a href="https://gsas.es/comprar-bases-de-datos-de-empresas/">Solicita Fichero Demo Excel gratuito </a>'
            ).replace(
                "www.pinkstone.es", '<a href="http://www.pinkstone.es/">www.pinkstone.es</a>'
            )
            html += f"<p>{para}</p>"

            # Insertar la tabla después de un párrafo específico
            if "Los campos que incluyen todas nuestras bases de datos son:" in para and tabla:
                html += "<table border='1'>"
                for fila in tabla:
                    html += "<tr>"
                    for celda in fila:
                        html += f"<td>{celda}</td>"
                    html += "</tr>"
                html += "</table>"
        except Exception as e:
         print(f"Error al convertir el contenido a HTML: {e}")
        html += "</body></html>"
        return html

    # Lista los archivos de Word en la carpeta y sus subcarpetas, ignorando los archivos temporales
    archivos_word = listar_archivos_docx(carpeta_word)

    if not archivos_word:
        print(f"No se encontraron archivos .docx en la ruta: {carpeta_word}")
    else:
        print(f"Se encontraron {len(archivos_word)} archivos .docx para procesar.")
        print("\nOrden de los archivos Word (por número):")
        for idx, archivo in enumerate(archivos_word, 1):
            nombre_archivo = os.path.basename(archivo)
            print(f"  {idx}. {nombre_archivo}")
        print()

    # Abre el archivo CSV y lee las URLs
    csv_path = os.path.join(carpeta_word, 'urls.csv')  # Asegúrate de que la ruta sea correcta
    try:
        # Primero lee todas las URLs para verificar el orden
        urls_list = []
        with open(csv_path, newline='', encoding='utf-8') as csvfile:
            url_reader = csv.reader(csvfile)
            for row in url_reader:
                if row and row[0].strip():  # Ignora filas vacías
                    urls_list.append(row[0].strip())
        
        print(f"Se encontraron {len(urls_list)} URLs en el CSV.")
        print("\nOrden de las URLs en el CSV:")
        for idx, url in enumerate(urls_list, 1):
            print(f"  Fila {idx}: {url}")
        print()
        
        # Verificar que el número de archivos coincida con el número de URLs
        if len(archivos_word) != len(urls_list):
            print(f"ADVERTENCIA: El número de archivos Word ({len(archivos_word)}) no coincide con el número de URLs ({len(urls_list)}).")
            print("Se procesarán solo los pares que coincidan.\n")
        
        # Ahora procesa cada par URL-Archivo
        with open(csv_path, newline='', encoding='utf-8') as csvfile:
            url_reader = csv.reader(csvfile)

            for i, row in enumerate(url_reader):
                if i >= len(archivos_word):
                    print("No hay más archivos .docx para asignar a las URLs.")
                    break
                if not row:
                    print(f"Fila {i+1} vacía. Saltando...")
                    continue
                product_url = row[0].strip()
                if not product_url:
                    print(f"URL en la fila {i+1} está vacía. Saltando...")
                    continue

                archivo = archivos_word[i]
                nombre_archivo = os.path.basename(archivo)
                print(f"\n{'='*80}")
                print(f"Procesando Fila {i+1}:")
                print(f"  URL: {product_url}")
                print(f"  Archivo Word: {nombre_archivo}")
                print(f"{'='*80}")

                driver.get(product_url)

                    # Intentar hacer clic en "Editar producto"
                try:
                        edit_product_button = WebDriverWait(driver, 60).until(
                            EC.element_to_be_clickable((By.LINK_TEXT, 'Editar producto'))
                        )
                        edit_product_button.click()
                except TimeoutException:
                        print(f"No se encontró el botón 'Editar producto' en la URL: {product_url}. Saltando...")
                        continue
                except Exception as e:
                        print(f"Error al hacer clic en 'Editar producto' en la URL: {product_url}. Error: {e}. Saltando...")
                        continue

                    # Extraer y procesar datos del archivo Word
                contenido = leer_documento(archivo)
                titulo = extraer_titulo(contenido)
                contenido_descripcion = extraer_descripcion(contenido)
                doc = Document(archivo)
                tabla = extraer_tabla(doc)
                descripcion_html = convertir_a_html(contenido_descripcion, tabla)
                contenido_corto = descripcion_corta(doc)
                etiq_prod = etiqueta_product(doc)
                tit_seo = titulo_seo(doc)
                meta_desc = meta_description(doc)
                frase_objetivo = frase_clave(doc)
                precio = extraer_precio(doc)

                print(f"Título extraído: {titulo}")
                print(f"Descripción en HTML (primeros 100 caracteres): {descripcion_html[:100]}...")

                    # Actualizar campos en WordPress
                try:
                        # Actualizar título
                        title_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'title'))
                        )
                        title_field.clear()
                        title_field.send_keys(titulo)
                        print(f"Título enviado a WordPress: {titulo}")

                        # Actualizar descripción principal
                        iframe = WebDriverWait(driver, 60).until(
                            EC.presence_of_element_located((By.ID, 'content_ifr'))
                        )
                        driver.switch_to.frame(iframe)

                        editor_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'tinymce'))
                        )
                        editor_field.clear()
                        driver.execute_script("arguments[0].innerHTML = arguments[1];", editor_field, descripcion_html)
                        driver.switch_to.default_content()

                        # Actualizar frase clave objetivo
                        frase_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'focus-keyword-input-metabox'))
                        )
                        frase_field.clear()
                        frase_field.send_keys(Keys.CONTROL, 'a')
                        frase_field.send_keys(Keys.DELETE)
                        frase_field.send_keys(frase_objetivo)
                        print("Frase actualizada: "+ frase_objetivo)

                        # Actualizar etiqueta de producto
                        etiq_prodfield = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'new-tag-product_tag'))
                        )
                        etiq_prodfield.clear()
                        etiq_prodfield.send_keys(etiq_prod)
                        etiq_prodfield.send_keys(Keys.RETURN)
                        print("Etiqueta actualizada actualizada: " + etiq_prod)

                        # Actualizar título SEO
                        titseo_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'yoast-google-preview-title-metabox'))
                        )
    
                        titseo_field.send_keys(Keys.CONTROL, 'a')
                        titseo_field.send_keys(tit_seo)
                        print("SEO actualizado: " + tit_seo)
                        # Actualizar meta descripción
                        metdesc_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, 'yoast-google-preview-description-metabox'))
                        )

                        metdesc_field.send_keys(Keys.CONTROL, 'a')
                        metdesc_field.send_keys(meta_desc)
                        print("Meta actualizado: " + meta_desc )
                        # Actualizar precio normal
                        price_field = WebDriverWait(driver, 60).until(
                            EC.visibility_of_element_located((By.ID, '_regular_price'))
                        )
                        price_field.clear()
                        price_field.send_keys(precio)
                        print("Precio actualizado: " + precio)
                        # Actualizar descripción corta
                        WebDriverWait(driver, 60).until(
                            EC.frame_to_be_available_and_switch_to_it((By.ID, "excerpt_ifr"))
                        )
                        short_description_body = WebDriverWait(driver, 20).until(
                            EC.visibility_of_element_located((By.ID, "tinymce"))
                        )
                        short_description_body.clear()
                        short_description_body.send_keys(contenido_corto)
                        print("Descripcion actualizada: " + contenido_corto)
                         # Volver al contexto principal
                        driver.switch_to.default_content()
                        actualizar_button = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.ID, "publish")))
                        driver.execute_script("arguments[0].click();", actualizar_button)

                        #Hace click en el botón después de desplazarse
                        actualizar_button.click()
                        print("Se hizo clic en el botón de 'Actualizar'.")

                        # Esperar brevemente para asegurar que la actualización se complete
                        time.sleep(120)

                except TimeoutException:
                        print(f"Tiempo de espera agotado al intentar actualizar la URL: {product_url}.")
                except Exception as e:
                        print(f"Error al actualizar el producto en la URL: {product_url}. Error: {e}")

    except FileNotFoundError:
        print(f"No se encontró el archivo CSV en la ruta: {csv_path}")
    except Exception as e:
        print(f"Error al abrir o procesar el archivo CSV: {e}")

finally:
    # Espera antes de cerrar el navegador
    time.sleep(150)  
    # Cierra el navegador
    driver.quit()
    print("Navegador cerrado.")
